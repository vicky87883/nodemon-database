import os
import fitz  # PyMuPDF for PDF files
from docx import Document
from groq import Groq
from django.shortcuts import render
from django.http import JsonResponse
from django.views.generic import ListView
from .models import ResearchPaper
from .forms import ResearchPaperForm
import re
import json
import logging
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import hashlib
import traceback
from difflib import SequenceMatcher
from collections import defaultdict
import time

# Configure logging for production
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the Groq client with the API key
client = Groq()

@dataclass
class ExtractionResult:
    """Data class for extraction results with confidence scores"""
    title: str
    authors: str
    abstract: str
    introduction: str
    methodology: str
    results: str
    conclusion: str
    references: str
    content_type: str
    confidence_score: float
    extraction_method: str
    error_log: List[str]
    processing_time: float
    word_count: int
    section_confidence: Dict[str, float]

class ProductionError(Exception):
    """Custom exception for production errors"""
    pass

def process_pdf_enterprise(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade PDF processing with comprehensive error handling.
    """
    errors = []
    text = ""
    
    try:
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        for page_num in range(total_pages):
            try:
                page = doc[page_num]
                page_text = page.get_text("text")
                text += page_text + "\n"
            except Exception as e:
                error_msg = f"Error processing page {page_num + 1}: {str(e)}"
                errors.append(error_msg)
                logger.warning(error_msg)
                continue
        
        doc.close()
        
        if not text.strip():
            raise ProductionError("No text extracted from PDF")
            
    except Exception as e:
        error_msg = f"Critical PDF processing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)
    
    return text, errors

def process_docx_enterprise(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade DOCX processing with comprehensive error handling.
    """
    errors = []
    text = ""
    
    try:
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        if not text.strip():
            raise ProductionError("No text extracted from DOCX")
            
    except Exception as e:
        error_msg = f"Critical DOCX processing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)
    
    return text, errors

def extract_text_from_file_enterprise(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade text extraction with comprehensive error handling.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text, errors = process_pdf_enterprise(file_path)
        elif file_extension == '.docx':
            text, errors = process_docx_enterprise(file_path)
        else:
            raise ProductionError(f"Unsupported file format: {file_extension}")
        
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Text extraction completed in {processing_time:.2f} seconds")
        
        return text, errors
        
    except Exception as e:
        error_msg = f"Critical text extraction error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)

def extract_document_sections_enterprise(text: str) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Enterprise-grade document section extraction with confidence scoring.
    """
    start_time = datetime.now()
    errors = []
    
    sections = {
        'title': '',
        'authors': '',
        'abstract': '',
        'introduction': '',
        'methodology': '',
        'results': '',
        'conclusion': '',
        'references': '',
        'keywords': '',
        'acknowledgments': '',
        'appendix': ''
    }
    
    confidence_scores = {
        'title': 0.0,
        'authors': 0.0,
        'abstract': 0.0,
        'introduction': 0.0,
        'methodology': 0.0,
        'results': 0.0,
        'conclusion': 0.0,
        'references': 0.0,
        'keywords': 0.0,
        'acknowledgments': 0.0,
        'appendix': 0.0
    }
    
    try:
        # Clean and normalize text
        text = text.strip()
        lines = text.split('\n')
        
        # Advanced title extraction with multiple strategies
        title, title_confidence = extract_title_enterprise(lines)
        sections['title'] = title
        confidence_scores['title'] = title_confidence
        
        # Advanced author extraction
        authors, authors_confidence = extract_authors_enterprise(lines)
        sections['authors'] = authors
        confidence_scores['authors'] = authors_confidence
        
        # Advanced section extraction with confidence scoring
        extracted_sections, section_confidences = extract_sections_enterprise(lines)
        sections.update(extracted_sections)
        confidence_scores.update(section_confidences)
        
        # Extract keywords
        keywords, keywords_confidence = extract_keywords_enterprise(text)
        sections['keywords'] = keywords
        confidence_scores['keywords'] = keywords_confidence
        
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Enterprise section extraction completed in {processing_time:.2f} seconds")
        
        return sections, confidence_scores, errors
        
    except Exception as e:
        error_msg = f"Enterprise section extraction error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        return sections, confidence_scores, errors

def extract_title_enterprise(lines: List[str]) -> Tuple[str, float]:
    """
    Enterprise-grade title extraction with confidence scoring.
    """
    title_candidates = []
    
    for i, line in enumerate(lines[:10]):
        line = line.strip()
        if not line or len(line) < 5:
            continue
            
        # Calculate title likelihood score
        score = 0.0
        
        # Length check (titles are usually 10-200 characters)
        if 10 <= len(line) <= 200:
            score += 0.3
        
        # Capitalization check
        if any(char.isupper() for char in line[:10]):
            score += 0.2
        
        # Avoid common non-title words
        non_title_words = ['abstract', 'introduction', 'chapter', 'section', 'author', 'authors', 'university', 'department']
        if not any(word in line.lower() for word in non_title_words):
            score += 0.2
        
        # Position bonus (first few lines are more likely to be titles)
        position_bonus = max(0, (10 - i) / 10)
        score += position_bonus * 0.2
        
        # Check for academic title patterns
        academic_patterns = [
            r'^[A-Z][^.!?]{10,100}[.!?]?\s*$',
            r'^[A-Z][A-Z\s]{10,100}\s*$',
            r'^[A-Z][a-z\s]{10,100}:\s*[A-Z]'
        ]
        
        for pattern in academic_patterns:
            if re.match(pattern, line):
                score += 0.1
                break
        
        if score > 0.3:
            title_candidates.append((line, score))
    
    if title_candidates:
        # Return the highest scoring title
        best_title, best_score = max(title_candidates, key=lambda x: x[1])
        return best_title, min(best_score, 1.0)
    
    return "Untitled", 0.0

def extract_authors_enterprise(lines: List[str]) -> Tuple[str, float]:
    """
    Enterprise-grade author extraction with confidence scoring.
    """
    author_candidates = []
    
    for i, line in enumerate(lines[:20]):
        line = line.strip()
        if not line or len(line) > 300:
            continue
            
        score = 0.0
        
        # Check for author indicators
        author_indicators = ['by ', 'author', 'authors', 'et al', 'and ', ',']
        if any(indicator in line.lower() for indicator in author_indicators):
            score += 0.4
        
        # Check for email patterns (often near authors)
        if '@' in line and '.' in line:
            if i > 0 and lines[i-1].strip():
                author_candidates.append((lines[i-1].strip(), 0.8))
        
        # Check for name patterns
        name_patterns = [
            r'([A-Z][a-z]+ [A-Z][a-z]+(?:, [A-Z][a-z]+ [A-Z][a-z]+)*)',
            r'([A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+)',
            r'([A-Z][a-z]+ [A-Z][a-z]+ et al)'
        ]
        
        for pattern in name_patterns:
            matches = re.findall(pattern, line)
            if matches:
                score += 0.3
                break
        
        # Position bonus
        position_bonus = max(0, (20 - i) / 20)
        score += position_bonus * 0.2
        
        if score > 0.3:
            author_candidates.append((line, score))
    
    if author_candidates:
        best_author, best_score = max(author_candidates, key=lambda x: x[1])
        return best_author, min(best_score, 1.0)
    
    return "Unknown", 0.0

def extract_sections_enterprise(lines: List[str]) -> Tuple[Dict[str, str], Dict[str, float]]:
    """
    Enterprise-grade section extraction with confidence scoring.
    """
    sections = {}
    confidence_scores = {}
    
    section_patterns = {
        'abstract': [r'abstract', r'abstract:', r'^\s*abstract\s*$'],
        'introduction': [r'introduction', r'introduction:', r'^\s*introduction\s*$'],
        'methodology': [r'methodology', r'methods', r'method', r'methodology:', r'methods:'],
        'results': [r'results', r'results:', r'^\s*results\s*$'],
        'conclusion': [r'conclusion', r'conclusions', r'conclusion:', r'^\s*conclusion\s*$'],
        'references': [r'references', r'reference', r'bibliography', r'works cited', r'references:'],
        'acknowledgments': [r'acknowledgments', r'acknowledgement', r'acknowledgments:'],
        'appendix': [r'appendix', r'appendices', r'appendix:']
    }
    
    current_section = None
    section_content = []
    section_start_line = 0
    
    for line_num, line in enumerate(lines):
        line_lower = line.lower().strip()
        line_original = line.strip()
        
        # Check for section headers
        found_section = None
        max_confidence = 0.0
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line_lower, re.IGNORECASE):
                    # Calculate confidence based on pattern match and line characteristics
                    confidence = 0.5  # Base confidence
                    
                    # Short lines are more likely to be headers
                    if len(line_lower) < 50:
                        confidence += 0.3
                    
                    # Check for formatting (all caps, bold indicators, etc.)
                    if line_original.isupper():
                        confidence += 0.2
                    
                    if confidence > max_confidence:
                        max_confidence = confidence
                        found_section = section_name
        
        if found_section and max_confidence > 0.6:
            # Save previous section
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
                confidence_scores[current_section] = calculate_section_confidence_enterprise(section_content, line_num - section_start_line)
            
            current_section = found_section
            section_content = []
            section_start_line = line_num
        elif current_section and line_original:
            section_content.append(line_original)
    
    # Save last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content)
        confidence_scores[current_section] = calculate_section_confidence_enterprise(section_content, len(lines) - section_start_line)
    
    # Initialize missing sections
    for section_name in section_patterns.keys():
        if section_name not in sections:
            sections[section_name] = ''
            confidence_scores[section_name] = 0.0
    
    return sections, confidence_scores

def extract_keywords_enterprise(text: str) -> Tuple[str, float]:
    """
    Enterprise-grade keyword extraction with confidence scoring.
    """
    keywords = ""
    confidence = 0.0
    
    # Look for keyword sections
    keyword_patterns = [
        r'keywords?[:\s]+([^.\n]+)',
        r'key\s+words?[:\s]+([^.\n]+)',
        r'index\s+terms?[:\s]+([^.\n]+)'
    ]
    
    for pattern in keyword_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            keywords = matches[0].strip()
            confidence = 0.8
            break
    
    return keywords, confidence

def calculate_section_confidence_enterprise(content: List[str], length: int) -> float:
    """
    Calculate confidence score for a section based on content quality.
    """
    if not content:
        return 0.0
    
    confidence = 0.0
    
    # Length-based confidence
    if length > 100:
        confidence += 0.3
    elif length > 50:
        confidence += 0.2
    elif length > 20:
        confidence += 0.1
    
    # Content quality indicators
    text = ' '.join(content)
    
    # Check for academic language
    academic_words = ['research', 'study', 'analysis', 'method', 'result', 'conclusion', 'data', 'experiment']
    academic_word_count = sum(1 for word in academic_words if word in text.lower())
    confidence += min(academic_word_count * 0.1, 0.3)
    
    # Check for proper formatting
    if any(char.isdigit() for char in text):
        confidence += 0.1
    
    if any(char.isupper() for char in text):
        confidence += 0.1
    
    return min(confidence, 1.0)

def analyze_extracted_content_with_groq_enterprise(extracted_sections: Dict[str, str], confidence_scores: Dict[str, float]) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Enterprise-grade AI analysis with confidence scoring and error handling.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        # Prepare content for AI analysis with confidence indicators
        content_summary = f"""
        EXTRACTED CONTENT WITH CONFIDENCE SCORES:
        
        TITLE: {extracted_sections.get('title', 'Not found')} (Confidence: {confidence_scores.get('title', 0.0):.2f})
        AUTHORS: {extracted_sections.get('authors', 'Not found')} (Confidence: {confidence_scores.get('authors', 0.0):.2f})
        ABSTRACT: {extracted_sections.get('abstract', 'Not found')[:800]} (Confidence: {confidence_scores.get('abstract', 0.0):.2f})
        INTRODUCTION: {extracted_sections.get('introduction', 'Not found')[:800]} (Confidence: {confidence_scores.get('introduction', 0.0):.2f})
        METHODOLOGY: {extracted_sections.get('methodology', 'Not found')[:800]} (Confidence: {confidence_scores.get('methodology', 0.0):.2f})
        RESULTS: {extracted_sections.get('results', 'Not found')[:800]} (Confidence: {confidence_scores.get('results', 0.0):.2f})
        CONCLUSION: {extracted_sections.get('conclusion', 'Not found')[:800]} (Confidence: {confidence_scores.get('conclusion', 0.0):.2f})
        REFERENCES: {extracted_sections.get('references', 'Not found')[:800]} (Confidence: {confidence_scores.get('references', 0.0):.2f})
        KEYWORDS: {extracted_sections.get('keywords', 'Not found')} (Confidence: {confidence_scores.get('keywords', 0.0):.2f})
        """
        
        prompt = f"""
        You are an expert research paper analyzer for a production-grade plagiarism detection system. 
        Analyze the extracted content and provide accurate categorization with confidence scores.
        
        {content_summary}
        
        TASK: Provide analysis in this exact format:
        
        TITLE: [Provide the correct title]
        AUTHORS: [Provide the correct authors]
        ABSTRACT: [Provide the complete abstract]
        INTRODUCTION: [Provide the introduction content]
        METHODOLOGY: [Provide the methodology content]
        RESULTS: [Provide the results content]
        CONCLUSION: [Provide the conclusion content]
        REFERENCES: [Provide the references content]
        KEYWORDS: [Provide the keywords]
        CONTENT_TYPE: [Classify as RESEARCH_PAPER, THESIS, or JOURNAL]
        OVERALL_CONFIDENCE: [Provide overall confidence score 0.0-1.0]
        
        CRITICAL REQUIREMENTS:
        1. Maintain high accuracy for plagiarism detection
        2. Provide complete sections, not summaries
        3. Use extracted content when confidence is high
        4. Improve content when confidence is low
        5. Mark sections as "Not found" if truly missing
        6. Provide realistic confidence scores
        """
        
        # Send request to Groq with production settings
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.05,  # Very low for consistency
            max_tokens=4000,
            top_p=0.9,
            frequency_penalty=0.1
        )
        
        response_content = chat_completion.choices[0].message.content
        logger.info("AI analysis completed successfully")
        
        # Parse the AI response
        return parse_llm_analysis_enterprise(response_content, extracted_sections, confidence_scores)
        
    except Exception as e:
        error_msg = f"AI analysis error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        
        # Return extracted content as fallback
        return extracted_sections, confidence_scores, errors

def parse_llm_analysis_enterprise(response_content: str, extracted_sections: Dict[str, str], confidence_scores: Dict[str, float]) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Enterprise-grade parsing of LLM response with confidence scoring.
    """
    result = extracted_sections.copy()
    final_confidence = confidence_scores.copy()
    errors = []
    
    try:
        lines = response_content.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Parse section headers
            if line.upper().startswith('TITLE:'):
                title = line.replace('TITLE:', '').strip()
                if title and title.lower() not in ['not found', 'none', '']:
                    result['title'] = title
                    final_confidence['title'] = min(confidence_scores.get('title', 0.0) + 0.2, 1.0)
                    
            elif line.upper().startswith('AUTHORS:'):
                authors = line.replace('AUTHORS:', '').strip()
                if authors and authors.lower() not in ['not found', 'none', 'unknown', '']:
                    result['authors'] = authors
                    final_confidence['authors'] = min(confidence_scores.get('authors', 0.0) + 0.2, 1.0)
                    
            elif line.upper().startswith('ABSTRACT:'):
                current_section = 'abstract'
                section_content = [line.replace('ABSTRACT:', '').strip()]
                
            elif line.upper().startswith('INTRODUCTION:'):
                current_section = 'introduction'
                section_content = [line.replace('INTRODUCTION:', '').strip()]
                
            elif line.upper().startswith('METHODOLOGY:'):
                current_section = 'methodology'
                section_content = [line.replace('METHODOLOGY:', '').strip()]
                
            elif line.upper().startswith('RESULTS:'):
                current_section = 'results'
                section_content = [line.replace('RESULTS:', '').strip()]
                
            elif line.upper().startswith('CONCLUSION:'):
                current_section = 'conclusion'
                section_content = [line.replace('CONCLUSION:', '').strip()]
                
            elif line.upper().startswith('REFERENCES:'):
                current_section = 'references'
                section_content = [line.replace('REFERENCES:', '').strip()]
                
            elif line.upper().startswith('KEYWORDS:'):
                keywords = line.replace('KEYWORDS:', '').strip()
                if keywords and keywords.lower() not in ['not found', 'none', '']:
                    result['keywords'] = keywords
                    final_confidence['keywords'] = min(confidence_scores.get('keywords', 0.0) + 0.2, 1.0)
                    
            elif line.upper().startswith('CONTENT_TYPE:'):
                content_type_raw = line.replace('CONTENT_TYPE:', '').strip().upper()
                if 'THESIS' in content_type_raw:
                    result['content_type'] = 'THESIS'
                elif 'JOURNAL' in content_type_raw:
                    result['content_type'] = 'JOURNAL'
                else:
                    result['content_type'] = 'RESEARCH_PAPER'
                    
            elif line.upper().startswith('OVERALL_CONFIDENCE:'):
                try:
                    overall_conf = float(line.replace('OVERALL_CONFIDENCE:', '').strip())
                    result['overall_confidence'] = min(max(overall_conf, 0.0), 1.0)
                except ValueError:
                    result['overall_confidence'] = 0.5
                    
            elif current_section and line:
                section_content.append(line)
                
            # Check for new section headers
            elif any(line.upper().startswith(header) for header in ['TITLE:', 'AUTHORS:', 'ABSTRACT:', 'INTRODUCTION:', 'METHODOLOGY:', 'RESULTS:', 'CONCLUSION:', 'REFERENCES:', 'KEYWORDS:', 'CONTENT_TYPE:', 'OVERALL_CONFIDENCE:']):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content)
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.1, 1.0)
                current_section = None
                section_content = []
        
        # Save last section
        if current_section and section_content:
            result[current_section] = ' '.join(section_content)
            final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.1, 1.0)
            
    except Exception as e:
        error_msg = f"LLM response parsing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
    
    return result, final_confidence, errors

def calculate_overall_confidence_enterprise(confidence_scores: Dict[str, float]) -> float:
    """
    Calculate overall confidence score for the entire extraction.
    """
    if not confidence_scores:
        return 0.0
    
    # Weight different sections based on importance for plagiarism detection
    weights = {
        'title': 0.1,
        'authors': 0.1,
        'abstract': 0.2,
        'introduction': 0.15,
        'methodology': 0.15,
        'results': 0.15,
        'conclusion': 0.1,
        'references': 0.05
    }
    
    weighted_sum = 0.0
    total_weight = 0.0
    
    for section, weight in weights.items():
        if section in confidence_scores:
            weighted_sum += confidence_scores[section] * weight
            total_weight += weight
    
    if total_weight > 0:
        return weighted_sum / total_weight
    
    return sum(confidence_scores.values()) / len(confidence_scores) if confidence_scores else 0.0

def analyze_content_with_groq_enterprise(text: str) -> ExtractionResult:
    """
    Enterprise-grade content analysis with comprehensive error handling and confidence scoring.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        logger.info("Starting enterprise-grade content analysis")
        
        # Step 1: Advanced document section extraction
        logger.info("Step 1: Advanced section extraction")
        extracted_sections, confidence_scores, extraction_errors = extract_document_sections_enterprise(text)
        errors.extend(extraction_errors)
        
        # Step 2: AI-powered analysis and enhancement
        logger.info("Step 2: AI-powered analysis")
        enhanced_sections, enhanced_confidence, ai_errors = analyze_extracted_content_with_groq_enterprise(extracted_sections, confidence_scores)
        errors.extend(ai_errors)
        
        # Calculate overall confidence
        overall_confidence = calculate_overall_confidence_enterprise(enhanced_confidence)
        
        # Calculate word count
        word_count = len(text.split())
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Determine extraction method
        extraction_method = "AI-Enhanced" if overall_confidence > 0.7 else "Traditional-Fallback"
        
        logger.info(f"Analysis completed in {processing_time:.2f} seconds with {overall_confidence:.2f} confidence")
        
        return ExtractionResult(
            title=enhanced_sections.get('title', 'Untitled'),
            authors=enhanced_sections.get('authors', 'Unknown'),
            abstract=enhanced_sections.get('abstract', 'No abstract found'),
            introduction=enhanced_sections.get('introduction', ''),
            methodology=enhanced_sections.get('methodology', ''),
            results=enhanced_sections.get('results', ''),
            conclusion=enhanced_sections.get('conclusion', ''),
            references=enhanced_sections.get('references', ''),
            content_type=enhanced_sections.get('content_type', 'RESEARCH_PAPER'),
            confidence_score=overall_confidence,
            extraction_method=extraction_method,
            error_log=errors,
            processing_time=processing_time,
            word_count=word_count,
            section_confidence=enhanced_confidence
        )
        
    except Exception as e:
        error_msg = f"Critical enterprise analysis error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        
        # Return minimal result with error information
        return ExtractionResult(
            title="Error in Analysis",
            authors="Unknown",
            abstract="Analysis failed",
            introduction="",
            methodology="",
            results="",
            conclusion="",
            references="",
            content_type="RESEARCH_PAPER",
            confidence_score=0.0,
            extraction_method="Error-Fallback",
            error_log=errors,
            processing_time=(datetime.now() - start_time).total_seconds(),
            word_count=0,
            section_confidence={}
        )

def upload_academic_content_enterprise(request):
    """
    Enterprise-grade file upload handler with comprehensive error handling and confidence reporting.
    """
    message = ''
    error = False
    title = ''
    confidence_info = {}
    
    if request.method == 'POST' and request.FILES.get('file'):
        form = ResearchPaperForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Save the file to the database
                paper = form.save()
                file_path = paper.file.path
                
                logger.info(f"Processing file: {file_path}")
                
                # Extract text from the uploaded file
                extracted_text, extraction_errors = extract_text_from_file_enterprise(file_path)
                
                if not extracted_text:
                    message = "Could not extract text from the uploaded file."
                    error = True
                    return render(request, 'upload.html', {
                        'form': form, 
                        'message': message, 
                        'error': error, 
                        'title': title
                    })
                
                # Perform enterprise-grade content analysis
                extraction_result = analyze_content_with_groq_enterprise(extracted_text)
                
                # Save extracted content into the database
                paper.title = extraction_result.title
                paper.authors = extraction_result.authors
                paper.abstract = extraction_result.abstract
                paper.introduction = extraction_result.introduction
                paper.methodology = extraction_result.methodology
                paper.results = extraction_result.results
                paper.conclusion = extraction_result.conclusion
                paper.references = extraction_result.references
                paper.content_type = extraction_result.content_type
                paper.save()
                
                # Prepare confidence information for display
                confidence_info = {
                    'overall_confidence': f"{extraction_result.confidence_score:.1%}",
                    'extraction_method': extraction_result.extraction_method,
                    'processing_time': f"{extraction_result.processing_time:.2f}s",
                    'word_count': extraction_result.word_count,
                    'section_confidence': extraction_result.section_confidence
                }
                
                message = f"File uploaded and processed successfully! Title: {extraction_result.title}"
                title = extraction_result.title
                
                # Log any errors for monitoring
                if extraction_result.error_log:
                    logger.warning(f"Processing completed with {len(extraction_result.error_log)} warnings")
                    for error_msg in extraction_result.error_log:
                        logger.warning(error_msg)
                
                # Pass all extracted data to template
                context = {
                    'form': form, 
                    'message': message, 
                    'error': error, 
                    'title': extraction_result.title,
                    'content_type': extraction_result.content_type,
                    'extracted_date': paper.extracted_date,
                    'abstract': extraction_result.abstract,
                    'introduction': extraction_result.introduction,
                    'methodology': extraction_result.methodology,
                    'results': extraction_result.results,
                    'conclusion': extraction_result.conclusion,
                    'references': extraction_result.references,
                    'confidence_info': confidence_info
                }
                return render(request, 'upload.html', context)
                
            except ProductionError as e:
                message = f"Production error: {str(e)}"
                error = True
                logger.error(f"Production error in upload: {str(e)}")
            except Exception as e:
                message = f"Unexpected error: {str(e)}"
                error = True
                logger.error(f"Unexpected error in upload: {str(e)}")
                logger.error(traceback.format_exc())
        else:
            message = "There was an issue with the file upload."
            error = True
    else:
        form = ResearchPaperForm()
    
    return render(request, 'upload.html', {
        'form': form, 
        'message': message, 
        'error': error, 
        'title': title,
        'confidence_info': confidence_info
    })

class AcademicContentList(ListView):
    model = ResearchPaper
    template_name = 'academic_content_list.html'
    context_object_name = 'academic_contents' 

@dataclass
class CaptionMatch:
    """Data class for caption-to-paragraph matching results"""
    caption_text: str
    caption_type: str  # 'figure' or 'table'
    caption_label: str  # e.g., 'Figure 1', 'Table 2'
    matching_paragraphs: List[str]
    match_scores: List[float]
    proximity_scores: List[float]
    keyword_scores: List[float]
    overall_confidence: float
    processing_time: float

def extract_captions_from_text(text: str) -> Tuple[List[Dict], List[Dict]]:
    """
    Extract figure and table captions from research paper text.
    """
    figures = []
    tables = []
    
    # Split text into lines for processing
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # Figure caption patterns
        figure_patterns = [
            r'^Figure\s+(\d+)[:.\s]+(.+)$',
            r'^Fig\.?\s*(\d+)[:.\s]+(.+)$',
            r'^Fig\s+(\d+)[:.\s]+(.+)$',
            r'^FIGURE\s+(\d+)[:.\s]+(.+)$'
        ]
        
        # Table caption patterns
        table_patterns = [
            r'^Table\s+(\d+)[:.\s]+(.+)$',
            r'^Tbl\.?\s*(\d+)[:.\s]+(.+)$',
            r'^TABLE\s+(\d+)[:.\s]+(.+)$'
        ]
        
        # Check for figure captions
        for pattern in figure_patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                figure_num = match.group(1)
                caption_text = match.group(2).strip()
                figures.append({
                    'label': f'Figure {figure_num}',
                    'number': figure_num,
                    'caption': caption_text,
                    'line_number': i,
                    'full_text': line
                })
                break
        
        # Check for table captions
        for pattern in table_patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                table_num = match.group(1)
                caption_text = match.group(2).strip()
                tables.append({
                    'label': f'Table {table_num}',
                    'number': table_num,
                    'caption': caption_text,
                    'line_number': i,
                    'full_text': line
                })
                break
    
    return figures, tables

def find_matching_paragraphs_for_caption(caption_text: str, caption_label: str, full_text: str, caption_line: int) -> Tuple[List[str], List[float], List[float], List[float]]:
    """
    Find paragraphs that match or refer to a given caption.
    """
    paragraphs = []
    match_scores = []
    proximity_scores = []
    keyword_scores = []
    
    # Split text into paragraphs
    text_lines = full_text.split('\n')
    
    # Extract keywords from caption
    caption_keywords = extract_keywords_from_caption(caption_text)
    
    # Find paragraphs that might reference this caption
    current_paragraph = []
    paragraph_start_line = 0
    
    for i, line in enumerate(text_lines):
        line = line.strip()
        
        # Check if this line starts a new paragraph (empty line or section header)
        if not line or is_section_header(line):
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                if len(paragraph_text.strip()) > 50:  # Only consider substantial paragraphs
                    # Calculate match scores
                    match_score = calculate_semantic_similarity(caption_text, paragraph_text)
                    proximity_score = calculate_proximity_score(i, caption_line, len(text_lines))
                    keyword_score = calculate_keyword_overlap(caption_keywords, paragraph_text)
                    
                    # Only include paragraphs with reasonable match scores
                    if match_score > 0.1 or keyword_score > 0.2 or proximity_score > 0.3:
                        paragraphs.append(paragraph_text)
                        match_scores.append(match_score)
                        proximity_scores.append(proximity_score)
                        keyword_scores.append(keyword_score)
                
                current_paragraph = []
                paragraph_start_line = i
        
        elif line:
            current_paragraph.append(line)
    
    # Process the last paragraph
    if current_paragraph:
        paragraph_text = ' '.join(current_paragraph)
        if len(paragraph_text.strip()) > 50:
            match_score = calculate_semantic_similarity(caption_text, paragraph_text)
            proximity_score = calculate_proximity_score(len(text_lines), caption_line, len(text_lines))
            keyword_score = calculate_keyword_overlap(caption_keywords, paragraph_text)
            
            if match_score > 0.1 or keyword_score > 0.2 or proximity_score > 0.3:
                paragraphs.append(paragraph_text)
                match_scores.append(match_score)
                proximity_scores.append(proximity_score)
                keyword_scores.append(keyword_score)
    
    return paragraphs, match_scores, proximity_scores, keyword_scores

def validate_caption_quality(caption: str, linked_text: str = None, full_text: str = None) -> Dict[str, float]:
    """
    Validate extracted figure and table captions for quality scoring.
    
    Args:
        caption (str): The caption to validate
        linked_text (str): Text that references the caption
        full_text (str): Full document text for broader context
        
    Returns:
        Dict[str, float]: Quality scores for clarity, contextual link, and completeness
    """
    if not caption or not caption.strip():
        return {
            "clarity_score": 0.0,
            "contextual_link_score": 0.0,
            "completeness_score": 0.0,
            "overall_quality_score": 0.0
        }
    
    # Calculate individual scores
    clarity_score = calculate_clarity_score(caption)
    contextual_link_score = calculate_contextual_link_score(caption, linked_text, full_text)
    completeness_score = calculate_completeness_score(caption, linked_text)
    
    # Calculate overall quality score (weighted average)
    overall_quality_score = (
        clarity_score * 0.3 +
        contextual_link_score * 0.4 +
        completeness_score * 0.3
    )
    
    return {
        "clarity_score": round(clarity_score, 3),
        "contextual_link_score": round(contextual_link_score, 3),
        "completeness_score": round(completeness_score, 3),
        "overall_quality_score": round(overall_quality_score, 3)
    }


def calculate_clarity_score(caption: str) -> float:
    """
    Calculate semantic clarity score for a caption.
    
    Args:
        caption (str): The caption to evaluate
        
    Returns:
        float: Clarity score between 0.0 and 1.0
    """
    if not caption or not caption.strip():
        return 0.0
    
    score = 1.0
    
    # Check for minimum length (captions should be descriptive)
    if len(caption.strip()) < 10:
        score -= 0.3
    elif len(caption.strip()) < 20:
        score -= 0.1
    
    # Check for common caption patterns
    caption_lower = caption.lower()
    
    # Penalize very generic captions
    generic_terms = ['figure', 'table', 'image', 'graph', 'chart', 'diagram']
    if any(term in caption_lower for term in generic_terms) and len(caption.split()) < 4:
        score -= 0.2
    
    # Bonus for descriptive words
    descriptive_words = ['showing', 'depicting', 'illustrating', 'demonstrating', 'comparing', 'analysis', 'results', 'performance', 'architecture', 'model', 'system']
    if any(word in caption_lower for word in descriptive_words):
        score += 0.1
    
    # Penalize for excessive punctuation or formatting issues
    if caption.count(':') > 2 or caption.count('.') > 3:
        score -= 0.1
    
    # Check for proper capitalization and structure
    words = caption.split()
    if len(words) > 0:
        # Check if first word is properly capitalized
        if not words[0][0].isupper():
            score -= 0.1
        
        # Check for mixed case issues
        all_upper = all(word.isupper() for word in words if len(word) > 1)
        if all_upper:
            score -= 0.1
    
    return max(0.0, min(1.0, score))


def calculate_contextual_link_score(caption: str, linked_text: str = None, full_text: str = None) -> float:
    """
    Calculate contextual link score based on how well the caption is referenced in text.
    
    Args:
        caption (str): The caption to evaluate
        linked_text (str): Text that references the caption
        full_text (str): Full document text for broader context
        
    Returns:
        float: Contextual link score between 0.0 and 1.0
    """
    if not caption or not caption.strip():
        return 0.0
    
    score = 0.0
    
    # Extract caption label (e.g., "Figure 3", "Table 1")
    caption_label = extract_caption_label(caption)
    
    if linked_text:
        # Check if caption label is mentioned in linked text
        if caption_label and caption_label.lower() in linked_text.lower():
            score += 0.6
        
        # Check for semantic similarity between caption and linked text
        caption_keywords = extract_keywords_from_caption(caption)
        text_keywords = extract_keywords_from_text(linked_text)
        
        # Calculate keyword overlap
        if caption_keywords and text_keywords:
            overlap = len(set(caption_keywords) & set(text_keywords))
            total_unique = len(set(caption_keywords) | set(text_keywords))
            if total_unique > 0:
                keyword_similarity = overlap / total_unique
                score += keyword_similarity * 0.3
        
        # Check for proximity and reference patterns
        if any(ref in linked_text.lower() for ref in ['as shown in', 'depicted in', 'illustrated in', 'presented in']):
            score += 0.1
    
    if full_text and not linked_text:
        # Search for caption label in full text
        if caption_label and caption_label.lower() in full_text.lower():
            score += 0.4
        
        # Check for semantic similarity with nearby text
        nearby_text = find_nearby_text_for_caption(caption, full_text)
        if nearby_text:
            caption_keywords = extract_keywords_from_caption(caption)
            text_keywords = extract_keywords_from_text(' '.join(nearby_text))
            
            if caption_keywords and text_keywords:
                overlap = len(set(caption_keywords) & set(text_keywords))
                total_unique = len(set(caption_keywords) | set(text_keywords))
                if total_unique > 0:
                    keyword_similarity = overlap / total_unique
                    score += keyword_similarity * 0.3
    
    return max(0.0, min(1.0, score))


def calculate_completeness_score(caption: str, linked_text: str = None) -> float:
    """
    Calculate completeness score based on how self-contained the caption is.
    
    Args:
        caption (str): The caption to evaluate
        linked_text (str): Text that references the caption
        
    Returns:
        float: Completeness score between 0.0 and 1.0
    """
    if not caption or not caption.strip():
        return 0.0
    
    score = 1.0
    
    # Check for essential caption components
    caption_lower = caption.lower()
    
    # Check if caption has a subject/object
    if not any(word in caption_lower for word in ['showing', 'of', 'for', 'with', 'in', 'on', 'at']):
        score -= 0.2
    
    # Check for descriptive content beyond just label
    words = caption.split()
    if len(words) < 3:  # Too short
        score -= 0.3
    elif len(words) < 5:  # Could be more descriptive
        score -= 0.1
    
    # Check for specific details (numbers, measurements, etc.)
    if re.search(r'\d+', caption):
        score += 0.1
    
    # Check for technical terms or domain-specific vocabulary
    technical_terms = ['accuracy', 'performance', 'model', 'system', 'architecture', 'algorithm', 'method', 'approach', 'technique', 'analysis', 'results', 'comparison', 'evaluation']
    if any(term in caption_lower for term in technical_terms):
        score += 0.1
    
    # Penalize for incomplete sentences or fragments
    if caption.endswith((':', '.', ';')) and len(words) < 4:
        score -= 0.2
    
    # Check if caption provides enough context without linked text
    if not linked_text:
        # Caption should be more self-contained
        if len(words) < 6:
            score -= 0.2
    
    return max(0.0, min(1.0, score))


def extract_caption_label(caption: str) -> str:
    """
    Extract the label part of a caption (e.g., "Figure 3", "Table 1").
    
    Args:
        caption (str): The caption text
        
    Returns:
        str: The extracted label or empty string
    """
    if not caption:
        return ""
    
    # Common patterns for figure/table labels
    patterns = [
        r'(Figure\s+\d+)',
        r'(Fig\.\s*\d+)',
        r'(Table\s+\d+)',
        r'(Tab\.\s*\d+)',
        r'(Figure\s+[A-Z])',
        r'(Table\s+[A-Z])'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, caption, re.IGNORECASE)
        if match:
            return match.group(1)
    
    return ""


def extract_keywords_from_caption(caption: str) -> List[str]:
    """
    Extract meaningful keywords from caption text.
    
    Args:
        caption (str): The caption text
        
    Returns:
        List[str]: List of keywords
    """
    if not caption:
        return []
    
    # Remove common stop words and caption labels
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'figure', 'fig', 'table', 'tab'}
    
    # Clean and tokenize
    cleaned = re.sub(r'[^\w\s]', ' ', caption.lower())
    words = cleaned.split()
    
    # Filter out stop words and short words
    keywords = [word for word in words if word not in stop_words and len(word) > 2]
    
    return keywords


def extract_keywords_from_text(text: str) -> List[str]:
    """
    Extract meaningful keywords from text.
    
    Args:
        text (str): The text to extract keywords from
        
    Returns:
        List[str]: List of keywords
    """
    if not text:
        return []
    
    # Remove common stop words
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'as', 'shown', 'in', 'figure', 'table'}
    
    # Clean and tokenize
    cleaned = re.sub(r'[^\w\s]', ' ', text.lower())
    words = cleaned.split()
    
    # Filter out stop words and short words
    keywords = [word for word in words if word not in stop_words and len(word) > 2]
    
    return keywords


def validate_captions_batch(captions_data: List[Dict]) -> List[Dict]:
    """
    Validate a batch of captions for quality scoring.
    
    Args:
        captions_data (List[Dict]): List of caption data with 'caption', 'linked_text', 'full_text'
        
    Returns:
        List[Dict]: List of validation results with quality scores
    """
    results = []
    
    for caption_data in captions_data:
        caption = caption_data.get('caption', '')
        linked_text = caption_data.get('linked_text', '')
        full_text = caption_data.get('full_text', '')
        
        quality_scores = validate_caption_quality(caption, linked_text, full_text)
        
        result = {
            'caption': caption,
            'linked_text': linked_text,
            'quality_scores': quality_scores,
            'validation_status': 'high_quality' if quality_scores['overall_quality_score'] >= 0.8 else 'medium_quality' if quality_scores['overall_quality_score'] >= 0.6 else 'low_quality'
        }
        
        results.append(result)
    
    return results


def process_caption_validation_request(request) -> JsonResponse:
    """
    API endpoint for caption quality validation.
    
    Expected POST data:
    {
        "captions": [
            {
                "caption": "Figure 3: Model architecture",
                "linked_text": "As shown in Figure 3, the model consists of three layers and a residual connection.",
                "full_text": "..." // optional
            }
        ]
    }
    """
    try:
        if request.method != 'POST':
            return JsonResponse({
                'error': 'Only POST method is supported',
                'status': 'error'
            }, status=405)
        
        data = json.loads(request.body)
        captions_data = data.get('captions', [])
        
        if not captions_data:
            return JsonResponse({
                'error': 'No captions provided',
                'status': 'error'
            }, status=400)
        
        start_time = time.time()
        
        # Validate captions
        validation_results = validate_captions_batch(captions_data)
        
        # Calculate summary statistics
        total_captions = len(validation_results)
        high_quality_count = sum(1 for result in validation_results if result['validation_status'] == 'high_quality')
        medium_quality_count = sum(1 for result in validation_results if result['validation_status'] == 'medium_quality')
        low_quality_count = sum(1 for result in validation_results if result['validation_status'] == 'low_quality')
        
        avg_clarity = sum(result['quality_scores']['clarity_score'] for result in validation_results) / total_captions
        avg_contextual = sum(result['quality_scores']['contextual_link_score'] for result in validation_results) / total_captions
        avg_completeness = sum(result['quality_scores']['completeness_score'] for result in validation_results) / total_captions
        avg_overall = sum(result['quality_scores']['overall_quality_score'] for result in validation_results) / total_captions
        
        result = {
            'validation_results': validation_results,
            'summary': {
                'total_captions': total_captions,
                'high_quality_count': high_quality_count,
                'medium_quality_count': medium_quality_count,
                'low_quality_count': low_quality_count,
                'average_scores': {
                    'clarity_score': round(avg_clarity, 3),
                    'contextual_link_score': round(avg_contextual, 3),
                    'completeness_score': round(avg_completeness, 3),
                    'overall_quality_score': round(avg_overall, 3)
                }
            },
            'processing_time': time.time() - start_time,
            'status': 'success'
        }
        
        return JsonResponse(result)
        
    except json.JSONDecodeError:
        return JsonResponse({
            'error': 'Invalid JSON data',
            'status': 'error'
        }, status=400)
    except Exception as e:
        logger.error(f"Caption validation error: {str(e)}")
        return JsonResponse({
            'error': f'Processing error: {str(e)}',
            'status': 'error'
        }, status=500)


def find_caption_matches(caption_text: str, full_text: str) -> CaptionMatch:
    """
    Main function to find matching paragraphs for a given caption.
    """
    start_time = datetime.now()
    
    # Extract caption information
    caption_type = 'figure' if 'figure' in caption_text.lower() or 'fig' in caption_text.lower() else 'table'
    
    # Extract caption label (e.g., "Figure 1", "Table 2")
    label_match = re.search(r'(Figure|Fig|Table|Tbl)\s*(\d+)', caption_text, re.IGNORECASE)
    caption_label = label_match.group(0) if label_match else "Unknown"
    
    # Find matching paragraphs using traditional methods
    paragraphs, match_scores, proximity_scores, keyword_scores = find_matching_paragraphs_for_caption(
        caption_text, caption_label, full_text, 0  # We don't have line number in this context
    )
    
    # Use AI to enhance the results
    ai_paragraphs, ai_confidences = analyze_caption_with_ai(caption_text, caption_label, paragraphs, full_text)
    
    # Combine traditional and AI results
    all_paragraphs = paragraphs + ai_paragraphs
    all_match_scores = match_scores + ai_confidences
    all_proximity_scores = proximity_scores + [0.5] * len(ai_paragraphs)  # Default proximity for AI results
    all_keyword_scores = keyword_scores + [0.5] * len(ai_paragraphs)  # Default keyword score for AI results
    
    # Calculate overall confidence
    if all_match_scores:
        overall_confidence = sum(all_match_scores) / len(all_match_scores)
    else:
        overall_confidence = 0.0
    
    processing_time = (datetime.now() - start_time).total_seconds()
    
    return CaptionMatch(
        caption_text=caption_text,
        caption_type=caption_type,
        caption_label=caption_label,
        matching_paragraphs=all_paragraphs[:2],  # Limit to 2 paragraphs
        match_scores=all_match_scores[:2],
        proximity_scores=all_proximity_scores[:2],
        keyword_scores=all_keyword_scores[:2],
        overall_confidence=overall_confidence,
        processing_time=processing_time
    )

def process_caption_matching_request(request) -> JsonResponse:
    """
    Handle caption matching requests via API.
    """
    try:
        if request.method == 'POST':
            data = json.loads(request.body)
            caption_text = data.get('caption_text', '')
            full_text = data.get('full_text', '')
            
            if not caption_text or not full_text:
                return JsonResponse({
                    'error': 'Missing caption_text or full_text',
                    'success': False
                })
            
            # Find matches
            caption_match = find_caption_matches(caption_text, full_text)
            
            # Prepare response
            response_data = {
                'success': True,
                'caption_text': caption_match.caption_text,
                'caption_type': caption_match.caption_type,
                'caption_label': caption_match.caption_label,
                'matching_paragraphs': caption_match.matching_paragraphs,
                'match_scores': caption_match.match_scores,
                'proximity_scores': caption_match.proximity_scores,
                'keyword_scores': caption_match.keyword_scores,
                'overall_confidence': caption_match.overall_confidence,
                'processing_time': caption_match.processing_time
            }
            
            return JsonResponse(response_data)
        
        else:
            return JsonResponse({
                'error': 'Only POST method is supported',
                'success': False
            })
    
    except Exception as e:
        logger.error(f"Caption matching error: {str(e)}")
        return JsonResponse({
            'error': f'Processing error: {str(e)}',
            'success': False
        }) 

@dataclass
class ImagePlaceholder:
    """Data class for image placeholder information"""
    filename: str
    image_type: str  # 'figure' or 'table'
    nearby_text: List[str]
    context_snippets: List[str]
    suggested_caption: str
    confidence_score: float
    reasoning: str

def extract_image_placeholders_from_text(text: str) -> List[str]:
    """
    Extract image placeholder references from OCR text.
    """
    placeholders = []
    
    # Common image placeholder patterns
    placeholder_patterns = [
        r'image_\d+\.(jpg|jpeg|png|gif|bmp)',
        r'fig_\d+\.(jpg|jpeg|png|gif|bmp)',
        r'table_\d+\.(jpg|jpeg|png|gif|bmp)',
        r'img_\d+\.(jpg|jpeg|png|gif|bmp)',
        r'figure_\d+\.(jpg|jpeg|png|gif|bmp)',
        r'[A-Za-z]+_\d{3,}\.(jpg|jpeg|png|gif|bmp)'
    ]
    
    for pattern in placeholder_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                # Pattern with capture groups
                placeholder = match[0] if match[0] else match[1]
            else:
                placeholder = match
            placeholders.append(placeholder)
    
    return list(set(placeholders))  # Remove duplicates

def find_nearby_text_for_placeholder(placeholder: str, text: str, context_window: int = 200) -> List[str]:
    """
    Find text snippets near an image placeholder reference.
    """
    nearby_snippets = []
    
    # Split text into sentences
    sentences = re.split(r'[.!?]+', text)
    
    for i, sentence in enumerate(sentences):
        sentence = sentence.strip()
        if not sentence:
            continue
        
        # Check if placeholder is mentioned in this sentence
        if placeholder.lower() in sentence.lower():
            # Get surrounding context
            start_idx = max(0, i - 2)
            end_idx = min(len(sentences), i + 3)
            
            context = ' '.join(sentences[start_idx:end_idx])
            if len(context) > context_window:
                context = context[:context_window] + "..."
            
            nearby_snippets.append(context)
    
    return nearby_snippets

def analyze_image_context_with_ai(placeholder: str, nearby_text: List[str], full_text: str) -> Tuple[str, float, str]:
    """
    Use AI to analyze context and suggest a caption for the image placeholder.
    """
    try:
        # Prepare context for AI analysis
        context_text = '\n'.join(nearby_text[:3])  # Use top 3 most relevant snippets
        
        prompt = f"""
        You are an expert research paper analyst. Analyze the context around an image placeholder and suggest an appropriate caption.

        IMAGE PLACEHOLDER: {placeholder}
        
        NEARBY TEXT CONTEXT:
        {context_text}
        
        FULL DOCUMENT CONTEXT (first 1000 characters):
        {full_text[:1000]}
        
        TASK: Based on the nearby text references and document context, suggest a likely description for this image.
        
        CONSIDER:
        1. What type of content is being discussed nearby?
        2. What would this image likely show based on the context?
        3. What is the academic purpose of this visual element?
        4. What terminology and concepts are mentioned in the vicinity?
        
        RESPONSE FORMAT:
        SUGGESTED_CAPTION: [Your suggested caption]
        CONFIDENCE: [Confidence score 0.0-1.0]
        REASONING: [Brief explanation of your reasoning]
        
        EXAMPLES:
        - If nearby text mentions "accuracy comparison", suggest "Accuracy comparison over training epochs"
        - If nearby text mentions "model architecture", suggest "Model architecture diagram"
        - If nearby text mentions "experimental results", suggest "Experimental results summary"
        """
        
        # Send to Groq for analysis
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.2,
            max_tokens=500,
            top_p=0.9
        )
        
        response = chat_completion.choices[0].message.content
        
        # Parse AI response
        suggested_caption = "Unknown image content"
        confidence_score = 0.5
        reasoning = "No specific context found"
        
        lines = response.split('\n')
        for line in lines:
            line = line.strip()
            if line.upper().startswith('SUGGESTED_CAPTION:'):
                suggested_caption = line.replace('SUGGESTED_CAPTION:', '').strip()
            elif line.upper().startswith('CONFIDENCE:'):
                try:
                    confidence_score = float(line.replace('CONFIDENCE:', '').strip())
                    confidence_score = min(max(confidence_score, 0.0), 1.0)
                except ValueError:
                    confidence_score = 0.5
            elif line.upper().startswith('REASONING:'):
                reasoning = line.replace('REASONING:', '').strip()
        
        return suggested_caption, confidence_score, reasoning
        
    except Exception as e:
        logger.error(f"AI image context analysis error: {str(e)}")
        return "Unknown image content", 0.3, f"Analysis failed: {str(e)}"

def determine_image_type(placeholder: str, nearby_text: List[str]) -> str:
    """
    Determine if the placeholder is likely a figure or table based on context.
    """
    # Check filename patterns
    if 'table' in placeholder.lower() or 'tbl' in placeholder.lower():
        return 'table'
    elif 'fig' in placeholder.lower() or 'figure' in placeholder.lower():
        return 'figure'
    
    # Check nearby text for clues
    text_content = ' '.join(nearby_text).lower()
    
    # Table indicators
    table_indicators = ['table', 'tabular', 'data', 'results', 'comparison', 'summary', 'statistics', 'values', 'columns', 'rows']
    table_score = sum(1 for indicator in table_indicators if indicator in text_content)
    
    # Figure indicators
    figure_indicators = ['figure', 'diagram', 'chart', 'graph', 'plot', 'visualization', 'image', 'photo', 'illustration', 'architecture']
    figure_score = sum(1 for indicator in figure_indicators if indicator in text_content)
    
    if table_score > figure_score:
        return 'table'
    elif figure_score > table_score:
        return 'figure'
    else:
        # Default to figure if unclear
        return 'figure'

def guess_image_descriptions(text: str, figure_placeholders: List[str] = None) -> Dict[str, str]:
    """
    Main function to guess descriptions for image placeholders.
    """
    start_time = datetime.now()
    results = {}
    
    try:
        # Extract placeholders if not provided
        if figure_placeholders is None:
            figure_placeholders = extract_image_placeholders_from_text(text)
        
        logger.info(f"Found {len(figure_placeholders)} image placeholders")
        
        for placeholder in figure_placeholders:
            # Find nearby text context
            nearby_text = find_nearby_text_for_placeholder(placeholder, text)
            
            if not nearby_text:
                # If no nearby text found, use general document context
                nearby_text = [text[:500] + "..." if len(text) > 500 else text]
            
            # Determine image type
            image_type = determine_image_type(placeholder, nearby_text)
            
            # Use AI to suggest caption
            suggested_caption, confidence, reasoning = analyze_image_context_with_ai(
                placeholder, nearby_text, text
            )
            
            # Store result
            results[placeholder] = suggested_caption
            
            # Log the analysis
            logger.info(f"Placeholder: {placeholder}")
            logger.info(f"Type: {image_type}")
            logger.info(f"Suggested Caption: {suggested_caption}")
            logger.info(f"Confidence: {confidence:.2f}")
            logger.info(f"Reasoning: {reasoning}")
            logger.info("-" * 50)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Image description guessing completed in {processing_time:.2f} seconds")
        
    except Exception as e:
        logger.error(f"Image description guessing error: {str(e)}")
        # Return empty results on error
        results = {}
    
    return results

def process_image_description_request(request) -> JsonResponse:
    """
    Handle image description guessing requests via API.
    """
    try:
        if request.method == 'POST':
            data = json.loads(request.body)
            figure_placeholders = data.get('figure_placeholders', [])
            text_snippets = data.get('text_snippets', [])
            
            if not figure_placeholders:
                return JsonResponse({
                    'error': 'No figure placeholders provided',
                    'success': False
                })
            
            # Combine text snippets into full text
            full_text = ' '.join(text_snippets) if text_snippets else ""
            
            if not full_text:
                return JsonResponse({
                    'error': 'No text content provided for analysis',
                    'success': False
                })
            
            # Guess descriptions
            descriptions = guess_image_descriptions(full_text, figure_placeholders)
            
            # Prepare response
            response_data = {
                'success': True,
                'descriptions': descriptions,
                'processing_time': 0.0,  # Will be calculated in the function
                'total_placeholders': len(figure_placeholders)
            }
            
            return JsonResponse(response_data)
        
        else:
            return JsonResponse({
                'error': 'Only POST method is supported',
                'success': False
            })
    
    except Exception as e:
        logger.error(f"Image description request error: {str(e)}")
        return JsonResponse({
            'error': f'Processing error: {str(e)}',
            'success': False
        })

def enhance_ocr_text_with_image_descriptions(ocr_text: str) -> Tuple[str, Dict[str, str]]:
    """
    Enhance OCR text by adding guessed descriptions for image placeholders.
    """
    # Extract image placeholders
    placeholders = extract_image_placeholders_from_text(ocr_text)
    
    if not placeholders:
        return ocr_text, {}
    
    # Guess descriptions
    descriptions = guess_image_descriptions(ocr_text, placeholders)
    
    # Enhance text by adding descriptions
    enhanced_text = ocr_text
    
    for placeholder, description in descriptions.items():
        # Add description after the placeholder
        placeholder_pattern = re.escape(placeholder)
        replacement = f"{placeholder}\n[Caption: {description}]"
        enhanced_text = re.sub(placeholder_pattern, replacement, enhanced_text, flags=re.IGNORECASE)
    
    return enhanced_text, descriptions 

def normalize_caption_for_fingerprinting(caption: str) -> str:
    """
    Normalize academic figure/table captions for fingerprinting and semantic comparison.
    
    Args:
        caption (str): Raw caption text
        
    Returns:
        str: Normalized caption ready for fingerprinting
    """
    if not caption or not caption.strip():
        return ""
    
    # Convert to lowercase
    normalized = caption.lower().strip()
    
    # Remove citations in parentheses (e.g., "(Smith, 2020)", "(et al., 2021)")
    normalized = re.sub(r'\([^)]*\)', '', normalized)
    
    # Remove citations in brackets (e.g., "[1]", "[2-5]")
    normalized = re.sub(r'\[[^\]]*\]', '', normalized)
    
    # Remove special characters except basic punctuation
    normalized = re.sub(r'[^\w\s\.\,\:\;\-\_]', '', normalized)
    
    # Replace numbers with [#] placeholder
    normalized = re.sub(r'\d+', '[#]', normalized)
    
    # Remove extra whitespace and normalize spaces
    normalized = re.sub(r'\s+', ' ', normalized)
    
    # Remove leading/trailing whitespace
    normalized = normalized.strip()
    
    return normalized


def normalize_captions_batch(captions: List[str]) -> List[str]:
    """
    Normalize a batch of captions for fingerprinting and semantic comparison.
    
    Args:
        captions (List[str]): List of raw caption texts
        
    Returns:
        List[str]: List of normalized captions
    """
    return [normalize_caption_for_fingerprinting(caption) for caption in captions]


def calculate_caption_similarity(caption1: str, caption2: str) -> float:
    """
    Calculate semantic similarity between two normalized captions.
    
    Args:
        caption1 (str): First caption
        caption2 (str): Second caption
        
    Returns:
        float: Similarity score between 0.0 and 1.0
    """
    # Normalize both captions
    norm1 = normalize_caption_for_fingerprinting(caption1)
    norm2 = normalize_caption_for_fingerprinting(caption2)
    
    if not norm1 or not norm2:
        return 0.0
    
    # Calculate similarity using SequenceMatcher
    similarity = SequenceMatcher(None, norm1, norm2).ratio()
    
    return similarity


def find_similar_captions(target_caption: str, caption_list: List[str], threshold: float = 0.7) -> List[Tuple[str, float]]:
    """
    Find captions similar to a target caption above a similarity threshold.
    
    Args:
        target_caption (str): The caption to find matches for
        caption_list (List[str]): List of captions to search through
        threshold (float): Minimum similarity score (0.0-1.0)
        
    Returns:
        List[Tuple[str, float]]: List of (caption, similarity_score) tuples
    """
    similar_captions = []
    
    for caption in caption_list:
        similarity = calculate_caption_similarity(target_caption, caption)
        if similarity >= threshold:
            similar_captions.append((caption, similarity))
    
    # Sort by similarity score (highest first)
    similar_captions.sort(key=lambda x: x[1], reverse=True)
    
    return similar_captions


def process_caption_normalization_request(request) -> JsonResponse:
    """
    API endpoint for caption normalization.
    
    Expected POST data:
    {
        "captions": ["Figure 2: Accuracy across 10 folds (Smith, 2020)", "Table 1. Summary of Results"],
        "operation": "normalize" | "similarity" | "find_similar",
        "target_caption": "Figure 2: Accuracy across 10 folds (Smith, 2020)",  # for similarity/find_similar
        "threshold": 0.7  # for find_similar
    }
    """
    try:
        if request.method != 'POST':
            return JsonResponse({
                'error': 'Only POST method is supported',
                'status': 'error'
            }, status=405)
        
        data = json.loads(request.body)
        captions = data.get('captions', [])
        operation = data.get('operation', 'normalize')
        
        if not captions:
            return JsonResponse({
                'error': 'No captions provided',
                'status': 'error'
            }, status=400)
        
        start_time = time.time()
        
        if operation == 'normalize':
            # Normalize captions
            normalized_captions = normalize_captions_batch(captions)
            
            result = {
                'operation': 'normalize',
                'original_captions': captions,
                'normalized_captions': normalized_captions,
                'processing_time': time.time() - start_time,
                'status': 'success'
            }
            
        elif operation == 'similarity':
            # Calculate similarity between two captions
            target_caption = data.get('target_caption')
            if not target_caption or len(captions) < 2:
                return JsonResponse({
                    'error': 'For similarity operation, provide target_caption and at least 2 captions',
                    'status': 'error'
                }, status=400)
            
            similarities = []
            for caption in captions:
                if caption != target_caption:
                    similarity = calculate_caption_similarity(target_caption, caption)
                    similarities.append({
                        'caption': caption,
                        'similarity_score': similarity
                    })
            
            # Sort by similarity score
            similarities.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            result = {
                'operation': 'similarity',
                'target_caption': target_caption,
                'comparison_captions': captions,
                'similarities': similarities,
                'processing_time': time.time() - start_time,
                'status': 'success'
            }
            
        elif operation == 'find_similar':
            # Find similar captions above threshold
            target_caption = data.get('target_caption')
            threshold = data.get('threshold', 0.7)
            
            if not target_caption:
                return JsonResponse({
                    'error': 'For find_similar operation, provide target_caption',
                    'status': 'error'
                }, status=400)
            
            similar_captions = find_similar_captions(target_caption, captions, threshold)
            
            result = {
                'operation': 'find_similar',
                'target_caption': target_caption,
                'threshold': threshold,
                'similar_captions': [
                    {'caption': caption, 'similarity_score': score}
                    for caption, score in similar_captions
                ],
                'total_matches': len(similar_captions),
                'processing_time': time.time() - start_time,
                'status': 'success'
            }
            
        else:
            return JsonResponse({
                'error': f'Unknown operation: {operation}. Supported: normalize, similarity, find_similar',
                'status': 'error'
            }, status=400)
        
        return JsonResponse(result)
        
    except json.JSONDecodeError:
        return JsonResponse({
            'error': 'Invalid JSON data',
            'status': 'error'
        }, status=400)
    except Exception as e:
        logger.error(f"Caption normalization error: {str(e)}")
        return JsonResponse({
            'error': f'Processing error: {str(e)}',
            'status': 'error'
        }, status=500) 

@dataclass
class OCRImageData:
    """Data structure for OCR-extracted image information"""
    image_filename: str
    text_snippets: List[str]
    bounding_box: Optional[Dict[str, float]] = None
    confidence: float = 0.0

@dataclass
class OCRProcessingResult:
    """Result of OCR processing with enhanced image descriptions"""
    full_text: str
    enhanced_text: str
    image_descriptions: Dict[str, str]
    image_confidence_scores: Dict[str, float]
    processing_summary: Dict[str, Any]

def process_ocr_document_with_images(
    full_text: str,
    ocr_images: List[str],
    ocr_text_snippets: List[str],
    groq_client=None
) -> OCRProcessingResult:
    """
    Process OCR-extracted document with image placeholders and text snippets.
    
    Args:
        full_text: Complete text extracted from PDF
        ocr_images: List of image filenames (e.g., ['image_001.jpg', 'image_002.jpg'])
        ocr_text_snippets: List of text snippets near image bounding boxes
        groq_client: Optional Groq client for AI analysis
    
    Returns:
        OCRProcessingResult with enhanced text and image descriptions
    """
    try:
        logger.info(f"Processing OCR document with {len(ocr_images)} images")
        
        # Initialize result structure
        image_descriptions = {}
        image_confidence_scores = {}
        enhanced_text = full_text
        
        # Process each image
        for i, image_filename in enumerate(ocr_images):
            text_snippet = ocr_text_snippets[i] if i < len(ocr_text_snippets) else ""
            
            # Generate description using AI if available
            if groq_client:
                description, confidence = analyze_ocr_image_with_ai(
                    image_filename, text_snippet, full_text, groq_client
                )
            else:
                description, confidence = analyze_ocr_image_traditional(
                    image_filename, text_snippet, full_text
                )
            
            image_descriptions[image_filename] = description
            image_confidence_scores[image_filename] = confidence
            
            # Enhance the text by adding the description
            enhanced_text = insert_image_description_into_text(
                enhanced_text, image_filename, description
            )
        
        # Calculate processing summary
        processing_summary = {
            "total_images": len(ocr_images),
            "processed_images": len(image_descriptions),
            "average_confidence": sum(image_confidence_scores.values()) / len(image_confidence_scores) if image_confidence_scores else 0.0,
            "enhancement_ratio": len(enhanced_text) / len(full_text) if full_text else 1.0
        }
        
        return OCRProcessingResult(
            full_text=full_text,
            enhanced_text=enhanced_text,
            image_descriptions=image_descriptions,
            image_confidence_scores=image_confidence_scores,
            processing_summary=processing_summary
        )
        
    except Exception as e:
        logger.error(f"Error processing OCR document: {str(e)}")
        return OCRProcessingResult(
            full_text=full_text,
            enhanced_text=full_text,
            image_descriptions={},
            image_confidence_scores={},
            processing_summary={"error": str(e)}
        )

def analyze_ocr_image_with_ai(
    image_filename: str,
    text_snippet: str,
    full_text: str,
    groq_client
) -> Tuple[str, float]:
    """
    Use AI to analyze OCR image and generate description.
    """
    try:
        # Determine image type from filename
        image_type = determine_image_type_from_filename(image_filename)
        
        # Find relevant context in full text
        context_snippets = find_relevant_context_for_image(
            image_filename, text_snippet, full_text
        )
        
        prompt = f"""
        You are analyzing an OCR-extracted academic document with image placeholders.
        
        Image: {image_filename}
        Image Type: {image_type}
        Nearby Text: {text_snippet}
        Context Snippets: {context_snippets[:3]}  # Top 3 most relevant
        
        Task: Generate a clear, academic description of what this image likely shows.
        
        Guidelines:
        - Be specific but concise (1-2 sentences)
        - Use academic language
        - Consider the image type ({image_type})
        - Reference the nearby text and context
        - Focus on the likely content/purpose
        
        Response format:
        DESCRIPTION: [Your description here]
        CONFIDENCE: [0.0-1.0 score]
        REASONING: [Brief explanation of your reasoning]
        """
        
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.3,
            max_tokens=300
        )
        
        response_content = response.choices[0].message.content
        
        # Parse response
        description = extract_description_from_ai_response(response_content)
        confidence = extract_confidence_from_ai_response(response_content)
        
        return description, confidence
        
    except Exception as e:
        logger.error(f"AI analysis failed for {image_filename}: {str(e)}")
        return f"Image placeholder: {image_filename}", 0.3

def analyze_ocr_image_traditional(
    image_filename: str,
    text_snippet: str,
    full_text: str
) -> Tuple[str, float]:
    """
    Traditional analysis without AI for OCR images.
    """
    try:
        image_type = determine_image_type_from_filename(image_filename)
        
        # Extract keywords from nearby text
        keywords = extract_keywords_from_text(text_snippet)
        
        # Find related content in full text
        related_content = find_related_content_in_text(
            keywords, full_text, context_window=200
        )
        
        # Generate description based on patterns
        description = generate_description_from_patterns(
            image_filename, image_type, keywords, related_content
        )
        
        # Calculate confidence based on keyword density and context
        confidence = calculate_description_confidence(
            keywords, related_content, text_snippet
        )
        
        return description, confidence
        
    except Exception as e:
        logger.error(f"Traditional analysis failed for {image_filename}: {str(e)}")
        return f"Image placeholder: {image_filename}", 0.2

def determine_image_type_from_filename(filename: str) -> str:
    """Determine if image is likely a figure, table, or other based on filename."""
    filename_lower = filename.lower()
    
    if any(keyword in filename_lower for keyword in ['fig', 'figure', 'diagram', 'chart']):
        return "figure"
    elif any(keyword in filename_lower for keyword in ['table', 'tab']):
        return "table"
    elif any(keyword in filename_lower for keyword in ['graph', 'plot']):
        return "graph"
    else:
        return "image"

def find_relevant_context_for_image(
    image_filename: str,
    text_snippet: str,
    full_text: str,
    max_snippets: int = 5
) -> List[str]:
    """Find relevant context snippets for an image in the full text."""
    try:
        # Extract keywords from text snippet
        snippet_keywords = extract_keywords_from_text(text_snippet)
        
        # Split full text into sentences
        sentences = split_text_into_sentences(full_text)
        
        # Score each sentence based on keyword overlap
        scored_sentences = []
        for sentence in sentences:
            sentence_keywords = extract_keywords_from_text(sentence)
            overlap = len(set(snippet_keywords) & set(sentence_keywords))
            score = overlap / len(snippet_keywords) if snippet_keywords else 0
            scored_sentences.append((sentence, score))
        
        # Sort by score and return top snippets
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        return [sentence for sentence, score in scored_sentences[:max_snippets] if score > 0]
        
    except Exception as e:
        logger.error(f"Error finding context for {image_filename}: {str(e)}")
        return []

def insert_image_description_into_text(
    text: str,
    image_filename: str,
    description: str
) -> str:
    """Insert image description into the text at appropriate location."""
    try:
        # Find the best location to insert the description
        # Look for references to the image filename
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if image_filename in line:
                # Insert description after the line containing the image reference
                lines.insert(i + 1, f"[{description}]")
                break
        else:
            # If no direct reference found, append to the end
            lines.append(f"\n[{image_filename}: {description}]")
        
        return '\n'.join(lines)
        
    except Exception as e:
        logger.error(f"Error inserting description for {image_filename}: {str(e)}")
        return text

def extract_description_from_ai_response(response: str) -> str:
    """Extract description from AI response."""
    try:
        if "DESCRIPTION:" in response:
            start = response.find("DESCRIPTION:") + len("DESCRIPTION:")
            end = response.find("CONFIDENCE:") if "CONFIDENCE:" in response else len(response)
            description = response[start:end].strip()
            return description if description else "Image placeholder"
        return "Image placeholder"
    except:
        return "Image placeholder"

def extract_confidence_from_ai_response(response: str) -> float:
    """Extract confidence score from AI response."""
    try:
        if "CONFIDENCE:" in response:
            start = response.find("CONFIDENCE:") + len("CONFIDENCE:")
            end = response.find("REASONING:") if "REASONING:" in response else len(response)
            confidence_text = response[start:end].strip()
            try:
                return float(confidence_text)
            except:
                return 0.5
        return 0.5
    except:
        return 0.5

def split_text_into_sentences(text: str) -> List[str]:
    """Split text into sentences using basic punctuation."""
    import re
    sentences = re.split(r'[.!?]+', text)
    return [s.strip() for s in sentences if s.strip()]

def find_related_content_in_text(
    keywords: List[str],
    text: str,
    context_window: int = 200
) -> str:
    """Find content related to keywords in the text."""
    try:
        if not keywords:
            return ""
        
        # Find positions of keywords in text
        positions = []
        for keyword in keywords:
            pos = text.lower().find(keyword.lower())
            if pos != -1:
                positions.append(pos)
        
        if not positions:
            return ""
        
        # Get text around the middle position
        middle_pos = sum(positions) // len(positions)
        start = max(0, middle_pos - context_window)
        end = min(len(text), middle_pos + context_window)
        
        return text[start:end]
        
    except Exception as e:
        logger.error(f"Error finding related content: {str(e)}")
        return ""

def generate_description_from_patterns(
    filename: str,
    image_type: str,
    keywords: List[str],
    related_content: str
) -> str:
    """Generate description using pattern matching."""
    try:
        if image_type == "figure":
            if any(word in related_content.lower() for word in ['architecture', 'model', 'system']):
                return f"System architecture diagram showing {', '.join(keywords[:3])}"
            elif any(word in related_content.lower() for word in ['performance', 'accuracy', 'results']):
                return f"Performance results chart displaying {', '.join(keywords[:3])}"
            else:
                return f"Figure showing {', '.join(keywords[:3]) if keywords else 'relevant data'}"
        
        elif image_type == "table":
            return f"Data table summarizing {', '.join(keywords[:3]) if keywords else 'key information'}"
        
        elif image_type == "graph":
            return f"Graph plotting {', '.join(keywords[:3]) if keywords else 'data trends'}"
        
        else:
            return f"Image containing {', '.join(keywords[:3]) if keywords else 'visual content'}"
            
    except Exception as e:
        logger.error(f"Error generating description: {str(e)}")
        return f"Image placeholder: {filename}"

def calculate_description_confidence(
    keywords: List[str],
    related_content: str,
    text_snippet: str
) -> float:
    """Calculate confidence score for generated description."""
    try:
        # Base confidence
        confidence = 0.3
        
        # Boost for keyword presence
        if keywords:
            confidence += 0.2
        
        # Boost for related content
        if related_content and len(related_content) > 50:
            confidence += 0.2
        
        # Boost for text snippet quality
        if text_snippet and len(text_snippet) > 20:
            confidence += 0.2
        
        # Boost for keyword overlap between snippet and content
        snippet_keywords = extract_keywords_from_text(text_snippet)
        content_keywords = extract_keywords_from_text(related_content)
        overlap = len(set(snippet_keywords) & set(content_keywords))
        if overlap > 0:
            confidence += 0.1
        
        return min(confidence, 1.0)
        
    except Exception as e:
        logger.error(f"Error calculating confidence: {str(e)}")
        return 0.3

def process_ocr_document_request(request):
    """
    API endpoint for processing OCR documents with images.
    """
    try:
        if request.method == 'POST':
            data = request.POST
            
            full_text = data.get('full_text', '')
            ocr_images = data.getlist('ocr_images') if hasattr(data, 'getlist') else []
            ocr_text_snippets = data.getlist('ocr_text_snippets') if hasattr(data, 'getlist') else []
            
            # Handle JSON input
            if not full_text and request.content_type == 'application/json':
                json_data = json.loads(request.body)
                full_text = json_data.get('full_text', '')
                ocr_images = json_data.get('ocr_images', [])
                ocr_text_snippets = json_data.get('ocr_text_snippets', [])
            
            if not full_text:
                return JsonResponse({
                    'error': 'full_text is required',
                    'status': 'error'
                }, status=400)
            
            # Initialize Groq client
            groq_client = None
            try:
                groq_client = Groq(api_key=os.getenv('GROQ_API_KEY'))
            except Exception as e:
                logger.warning(f"Groq client initialization failed: {str(e)}")
            
            # Process the OCR document
            result = process_ocr_document_with_images(
                full_text=full_text,
                ocr_images=ocr_images,
                ocr_text_snippets=ocr_text_snippets,
                groq_client=groq_client
            )
            
            return JsonResponse({
                'status': 'success',
                'enhanced_text': result.enhanced_text,
                'image_descriptions': result.image_descriptions,
                'image_confidence_scores': result.image_confidence_scores,
                'processing_summary': result.processing_summary
            })
        
        return JsonResponse({
            'error': 'Only POST method is supported',
            'status': 'error'
        }, status=405)
        
    except Exception as e:
        logger.error(f"Error processing OCR document request: {str(e)}")
        return JsonResponse({
            'error': str(e),
            'status': 'error'
        }, status=500) 