import os
import fitz  # PyMuPDF for PDF files
from docx import Document
from groq import Groq
from django.shortcuts import render
from django.http import JsonResponse
from django.views.generic import ListView
from .models import ResearchPaper
from .forms import ResearchPaperForm
import re
import json
import logging
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
from datetime import datetime
import hashlib
import traceback

# Configure logging for production
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the Groq client with the API key
client = Groq(api_key="")

@dataclass
class ExtractionResult:
    """Data class for extraction results with confidence scores"""
    title: str
    authors: str
    abstract: str
    introduction: str
    methodology: str
    results: str
    conclusion: str
    references: str
    content_type: str
    confidence_score: float
    extraction_method: str
    error_log: List[str]
    processing_time: float
    word_count: int
    section_confidence: Dict[str, float]

class ProductionError(Exception):
    """Custom exception for production errors"""
    pass

def process_pdf(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade PDF processing with comprehensive error handling.
    """
    errors = []
    text = ""
    
    try:
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        for page_num in range(total_pages):
            try:
                page = doc[page_num]
                page_text = page.get_text("text")
                text += page_text + "\n"
            except Exception as e:
                error_msg = f"Error processing page {page_num + 1}: {str(e)}"
                errors.append(error_msg)
                logger.warning(error_msg)
                continue
        
        doc.close()
        
        if not text.strip():
            raise ProductionError("No text extracted from PDF")
            
    except Exception as e:
        error_msg = f"Critical PDF processing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)
    
    return text, errors

def process_docx(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade DOCX processing with comprehensive error handling.
    """
    errors = []
    text = ""
    
    try:
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        if not text.strip():
            raise ProductionError("No text extracted from DOCX")
            
    except Exception as e:
        error_msg = f"Critical DOCX processing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)
    
    return text, errors

def process_tex(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade LaTeX processing with comprehensive error handling.
    """
    errors = []
        text = ""
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        if not text.strip():
            raise ProductionError("No text extracted from LaTeX")
            
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e:
            error_msg = f"Critical LaTeX encoding error: {str(e)}"
            errors.append(error_msg)
            logger.error(error_msg)
            raise ProductionError(error_msg)
    except Exception as e:
        error_msg = f"Critical LaTeX processing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)
    
    return text, errors

def extract_text_from_file(file_path: str) -> Tuple[str, List[str]]:
    """
    Enterprise-grade text extraction with comprehensive error handling.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text, errors = process_pdf(file_path)
        elif file_extension == '.docx':
            text, errors = process_docx(file_path)
        elif file_extension == '.tex':
            text, errors = process_tex(file_path)
        else:
            raise ProductionError(f"Unsupported file format: {file_extension}")
        
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Text extraction completed in {processing_time:.2f} seconds")
        
        return text, errors
        
    except Exception as e:
        error_msg = f"Critical text extraction error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        raise ProductionError(error_msg)

def extract_document_sections_advanced(text: str) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Advanced document section extraction with confidence scoring.
    """
    start_time = datetime.now()
    errors = []
    
    sections = {
        'title': '',
        'authors': '',
        'abstract': '',
        'introduction': '',
        'methodology': '',
        'results': '',
        'conclusion': '',
        'references': '',
        'keywords': '',
        'acknowledgments': '',
        'appendix': ''
    }
    
    confidence_scores = {
        'title': 0.0,
        'authors': 0.0,
        'abstract': 0.0,
        'introduction': 0.0,
        'methodology': 0.0,
        'results': 0.0,
        'conclusion': 0.0,
        'references': 0.0,
        'keywords': 0.0,
        'acknowledgments': 0.0,
        'appendix': 0.0
    }
    
    try:
        # Clean and normalize text
        text = text.strip()
        lines = text.split('\n')
        
        # Advanced title extraction with multiple strategies
        title, title_confidence = extract_title_advanced(lines)
        sections['title'] = title
        confidence_scores['title'] = title_confidence
        
        # Advanced author extraction
        authors, authors_confidence = extract_authors_advanced(lines)
        sections['authors'] = authors
        confidence_scores['authors'] = authors_confidence
        
        # Advanced section extraction with confidence scoring
        extracted_sections, section_confidences = extract_sections_advanced(lines)
        sections.update(extracted_sections)
        confidence_scores.update(section_confidences)
        
        # Extract keywords
        keywords, keywords_confidence = extract_keywords_advanced(text)
        sections['keywords'] = keywords
        confidence_scores['keywords'] = keywords_confidence
        
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Advanced section extraction completed in {processing_time:.2f} seconds")
        
        return sections, confidence_scores, errors
        
    except Exception as e:
        error_msg = f"Advanced section extraction error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        return sections, confidence_scores, errors

def extract_title_advanced(lines: List[str]) -> Tuple[str, float]:
    """
    Advanced title extraction with confidence scoring.
    """
    title_candidates = []
    
    for i, line in enumerate(lines[:10]):
        line = line.strip()
        if not line or len(line) < 5:
            continue
            
        # Calculate title likelihood score
        score = 0.0
        
        # Length check (titles are usually 10-200 characters)
        if 10 <= len(line) <= 200:
            score += 0.3
        
        # Capitalization check
        if any(char.isupper() for char in line[:10]):
            score += 0.2
        
        # Avoid common non-title words
        non_title_words = ['abstract', 'introduction', 'chapter', 'section', 'author', 'authors', 'university', 'department']
        if not any(word in line.lower() for word in non_title_words):
            score += 0.2
        
        # Position bonus (first few lines are more likely to be titles)
        position_bonus = max(0, (10 - i) / 10)
        score += position_bonus * 0.2
        
        # Check for academic title patterns
        academic_patterns = [
            r'^[A-Z][^.!?]{10,100}[.!?]?\s*$',
            r'^[A-Z][A-Z\s]{10,100}\s*$',
            r'^[A-Z][a-z\s]{10,100}:\s*[A-Z]'
        ]
        
        for pattern in academic_patterns:
            if re.match(pattern, line):
                score += 0.1
                break
        
        if score > 0.3:
            title_candidates.append((line, score))
    
    if title_candidates:
        # Return the highest scoring title
        best_title, best_score = max(title_candidates, key=lambda x: x[1])
        return best_title, min(best_score, 1.0)
    
    return "Untitled", 0.0

def extract_authors_advanced(lines: List[str]) -> Tuple[str, float]:
    """
    Advanced author extraction with confidence scoring.
    """
    author_candidates = []
    
    for i, line in enumerate(lines[:20]):
        line = line.strip()
        if not line or len(line) > 300:
            continue
            
        score = 0.0
        
        # Check for author indicators
        author_indicators = ['by ', 'author', 'authors', 'et al', 'and ', ',']
        if any(indicator in line.lower() for indicator in author_indicators):
            score += 0.4
        
        # Check for email patterns (often near authors)
        if '@' in line and '.' in line:
            if i > 0 and lines[i-1].strip():
                author_candidates.append((lines[i-1].strip(), 0.8))
        
        # Check for name patterns
        name_patterns = [
            r'([A-Z][a-z]+ [A-Z][a-z]+(?:, [A-Z][a-z]+ [A-Z][a-z]+)*)',
            r'([A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+)',
            r'([A-Z][a-z]+ [A-Z][a-z]+ et al)'
        ]
        
        for pattern in name_patterns:
            matches = re.findall(pattern, line)
            if matches:
                score += 0.3
                break
        
        # Position bonus
        position_bonus = max(0, (20 - i) / 20)
        score += position_bonus * 0.2
        
        if score > 0.3:
            author_candidates.append((line, score))
    
    if author_candidates:
        best_author, best_score = max(author_candidates, key=lambda x: x[1])
        return best_author, min(best_score, 1.0)
    
    return "Unknown", 0.0

def extract_sections_advanced(lines: List[str]) -> Tuple[Dict[str, str], Dict[str, float]]:
    """
    Advanced section extraction with confidence scoring.
    """
    sections = {}
    confidence_scores = {}
    
    section_patterns = {
        'abstract': [r'abstract', r'abstract:', r'^\s*abstract\s*$'],
        'introduction': [r'introduction', r'introduction:', r'^\s*introduction\s*$'],
        'methodology': [r'methodology', r'methods', r'method', r'methodology:', r'methods:'],
        'results': [r'results', r'results:', r'^\s*results\s*$'],
        'conclusion': [r'conclusion', r'conclusions', r'conclusion:', r'^\s*conclusion\s*$'],
        'references': [r'references', r'reference', r'bibliography', r'works cited', r'references:'],
        'acknowledgments': [r'acknowledgments', r'acknowledgement', r'acknowledgments:'],
        'appendix': [r'appendix', r'appendices', r'appendix:']
    }
    
    current_section = None
    section_content = []
    section_start_line = 0
    
    for line_num, line in enumerate(lines):
        line_lower = line.lower().strip()
        line_original = line.strip()
        
        # Check for section headers
        found_section = None
        max_confidence = 0.0
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line_lower, re.IGNORECASE):
                    # Calculate confidence based on pattern match and line characteristics
                    confidence = 0.5  # Base confidence
                    
                    # Short lines are more likely to be headers
                    if len(line_lower) < 50:
                        confidence += 0.3
                    
                    # Check for formatting (all caps, bold indicators, etc.)
                    if line_original.isupper():
                        confidence += 0.2
                    
                    if confidence > max_confidence:
                        max_confidence = confidence
                        found_section = section_name
        
        if found_section and max_confidence > 0.6:
            # Save previous section
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
                confidence_scores[current_section] = calculate_section_confidence(section_content, line_num - section_start_line)
            
            current_section = found_section
            section_content = []
            section_start_line = line_num
        elif current_section and line_original:
            section_content.append(line_original)
    
    # Save last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content)
        confidence_scores[current_section] = calculate_section_confidence(section_content, len(lines) - section_start_line)
    
    # Initialize missing sections
    for section_name in section_patterns.keys():
        if section_name not in sections:
            sections[section_name] = ''
            confidence_scores[section_name] = 0.0
    
    return sections, confidence_scores

def extract_keywords_advanced(text: str) -> Tuple[str, float]:
    """
    Advanced keyword extraction with confidence scoring.
    """
    keywords = ""
    confidence = 0.0
    
    # Look for keyword sections
    keyword_patterns = [
        r'keywords?[:\s]+([^.\n]+)',
        r'key\s+words?[:\s]+([^.\n]+)',
        r'index\s+terms?[:\s]+([^.\n]+)'
    ]
    
    for pattern in keyword_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            keywords = matches[0].strip()
            confidence = 0.8
            break
    
    return keywords, confidence

def calculate_section_confidence(content: List[str], length: int) -> float:
    """
    Calculate confidence score for a section based on content quality.
    """
    if not content:
        return 0.0
    
    confidence = 0.0
    
    # Length-based confidence
    if length > 100:
        confidence += 0.3
    elif length > 50:
        confidence += 0.2
    elif length > 20:
        confidence += 0.1
    
    # Content quality indicators
    text = ' '.join(content)
    
    # Check for academic language
    academic_words = ['research', 'study', 'analysis', 'method', 'result', 'conclusion', 'data', 'experiment']
    academic_word_count = sum(1 for word in academic_words if word in text.lower())
    confidence += min(academic_word_count * 0.1, 0.3)
    
    # Check for proper formatting
    if any(char.isdigit() for char in text):
        confidence += 0.1
    
    if any(char.isupper() for char in text):
        confidence += 0.1
    
    return min(confidence, 1.0)

def analyze_extracted_content_with_groq_advanced(extracted_sections: Dict[str, str], confidence_scores: Dict[str, float]) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Advanced AI analysis with properly tuned prompts for accurate column-wise data insertion.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        # Create a structured summary of extracted content
        content_summary = f"""
        DOCUMENT CONTENT SUMMARY:
        
        TITLE: {extracted_sections.get('title', 'Not found')}
        AUTHORS: {extracted_sections.get('authors', 'Not found')}
        ABSTRACT: {extracted_sections.get('abstract', 'Not found')[:1000]}
        INTRODUCTION: {extracted_sections.get('introduction', 'Not found')[:1000]}
        METHODOLOGY: {extracted_sections.get('methodology', 'Not found')[:1000]}
        RESULTS: {extracted_sections.get('results', 'Not found')[:1000]}
        CONCLUSION: {extracted_sections.get('conclusion', 'Not found')[:1000]}
        REFERENCES: {extracted_sections.get('references', 'Not found')[:1000]}
        KEYWORDS: {extracted_sections.get('keywords', 'Not found')}
        """
        
        # Tuned prompt for accurate column-wise data extraction
        prompt = f"""
        You are a precision research paper analyzer for a production database system. 
        Your task is to extract and categorize content with 100% accuracy for database insertion.
        
        {content_summary}
        
        CRITICAL INSTRUCTIONS:
        1. Analyze the provided content carefully
        2. Extract ONLY the requested information for each field
        3. Provide complete, accurate content for database columns
        4. Use the exact format specified below
        5. If information is missing, mark as "Not found"
        6. Ensure all text is properly formatted and complete
        
        RESPONSE FORMAT (copy this exactly):
        
        TITLE: [Extract the complete paper title]
        AUTHORS: [Extract all author names in format: "Author1, Author2, Author3"]
        ABSTRACT: [Extract the complete abstract text]
        INTRODUCTION: [Extract the complete introduction section]
        METHODOLOGY: [Extract the complete methodology/methods section]
        RESULTS: [Extract the complete results section]
        CONCLUSION: [Extract the complete conclusion section]
        REFERENCES: [Extract the complete references/bibliography section]
        CONTENT_TYPE: [Classify as: RESEARCH_PAPER, THESIS, or JOURNAL]
        CONFIDENCE: [Provide confidence score 0.0-1.0]
        
        IMPORTANT:
        - Provide complete sections, not summaries
        - Maintain original formatting where possible
        - Ensure all text is properly extracted
        - Use "Not found" only if section is truly missing
        - Be precise and accurate for database insertion
        """
        
        # Send request to Groq with optimized settings for accuracy
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.01,  # Very low temperature for consistency
            max_tokens=6000,   # Increased for complete sections
            top_p=0.95,
            frequency_penalty=0.0,
            presence_penalty=0.0
        )
        
        response_content = chat_completion.choices[0].message.content
        logger.info("AI analysis completed successfully")
        
        # Parse the AI response with enhanced validation
        return parse_llm_response_enhanced(response_content, extracted_sections, confidence_scores)
        
    except Exception as e:
        error_msg = f"AI analysis error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        
        # Return extracted content as fallback
        return extracted_sections, confidence_scores, errors

def parse_llm_response_enhanced(response_content: str, extracted_sections: Dict[str, str], confidence_scores: Dict[str, float]) -> Tuple[Dict[str, str], Dict[str, float], List[str]]:
    """
    Enhanced parsing of LLM response with robust validation for database insertion.
    """
    result = extracted_sections.copy()
    final_confidence = confidence_scores.copy()
    errors = []
    
    try:
        # Clean and normalize the response
        response_content = response_content.strip()
        lines = response_content.split('\n')
        
        current_section = None
        section_content = []
        overall_confidence = 0.5  # Default confidence
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Parse section headers with exact matching
            if line.upper().startswith('TITLE:'):
                # Save previous section if exists
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                # Extract title
                title = line.replace('TITLE:', '').strip()
                if title and title.lower() not in ['not found', 'none', '']:
                    result['title'] = title
                    final_confidence['title'] = min(confidence_scores.get('title', 0.0) + 0.4, 1.0)
                current_section = None
                section_content = []
                
            elif line.upper().startswith('AUTHORS:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                # Extract authors
                authors = line.replace('AUTHORS:', '').strip()
                if authors and authors.lower() not in ['not found', 'none', 'unknown', '']:
                    result['authors'] = authors
                    final_confidence['authors'] = min(confidence_scores.get('authors', 0.0) + 0.4, 1.0)
                current_section = None
                section_content = []
                
            elif line.upper().startswith('ABSTRACT:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'abstract'
                section_content = [line.replace('ABSTRACT:', '').strip()]
                
            elif line.upper().startswith('INTRODUCTION:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'introduction'
                section_content = [line.replace('INTRODUCTION:', '').strip()]
                
            elif line.upper().startswith('METHODOLOGY:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'methodology'
                section_content = [line.replace('METHODOLOGY:', '').strip()]
                
            elif line.upper().startswith('RESULTS:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'results'
                section_content = [line.replace('RESULTS:', '').strip()]
                
            elif line.upper().startswith('CONCLUSION:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'conclusion'
                section_content = [line.replace('CONCLUSION:', '').strip()]
                
            elif line.upper().startswith('REFERENCES:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                current_section = 'references'
                section_content = [line.replace('REFERENCES:', '').strip()]
                
            elif line.upper().startswith('CONTENT_TYPE:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                content_type_raw = line.replace('CONTENT_TYPE:', '').strip().upper()
                if 'THESIS' in content_type_raw:
                    result['content_type'] = 'THESIS'
                elif 'JOURNAL' in content_type_raw:
                    result['content_type'] = 'JOURNAL'
        else:
                    result['content_type'] = 'RESEARCH_PAPER'
                current_section = None
                section_content = []
                
            elif line.upper().startswith('CONFIDENCE:'):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                
                try:
                    overall_confidence = float(line.replace('CONFIDENCE:', '').strip())
                    overall_confidence = min(max(overall_confidence, 0.0), 1.0)
                except ValueError:
                    overall_confidence = 0.5
                current_section = None
                section_content = []
                
            # Check for new section headers (to avoid content mixing)
            elif any(line.upper().startswith(header) for header in ['TITLE:', 'AUTHORS:', 'ABSTRACT:', 'INTRODUCTION:', 'METHODOLOGY:', 'RESULTS:', 'CONCLUSION:', 'REFERENCES:', 'CONTENT_TYPE:', 'CONFIDENCE:']):
                # Save previous section
                if current_section and section_content:
                    result[current_section] = ' '.join(section_content).strip()
                    final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
                current_section = None
                section_content = []
                
            elif current_section and line:
                # Add content to current section
                section_content.append(line)
        
        # Save the last section
        if current_section and section_content:
            result[current_section] = ' '.join(section_content).strip()
            final_confidence[current_section] = min(confidence_scores.get(current_section, 0.0) + 0.3, 1.0)
        
        # Validate and clean the results
        result = validate_and_clean_extracted_data(result)
        
        # Update overall confidence
        result['overall_confidence'] = overall_confidence
        
        logger.info("Enhanced LLM response parsing completed successfully")
        
        except Exception as e:
        error_msg = f"Enhanced LLM response parsing error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
    
    return result, final_confidence, errors

def validate_and_clean_extracted_data(data: Dict[str, str]) -> Dict[str, str]:
    """
    Validate and clean extracted data for database insertion.
    """
    cleaned_data = data.copy()
    
    # Clean and validate each field
    for field, value in cleaned_data.items():
        if isinstance(value, str):
            # Remove excessive whitespace
            value = ' '.join(value.split())
            
            # Remove common artifacts
            value = value.replace('Not found', '').replace('None', '').replace('N/A', '')
            
            # Ensure minimum length for important fields
            if field in ['title', 'authors'] and len(value.strip()) < 3:
                value = 'Unknown' if field == 'authors' else 'Untitled'
            
            # Truncate if too long for database constraints
            if field == 'title' and len(value) > 255:
                value = value[:252] + '...'
            elif field == 'authors' and len(value) > 500:
                value = value[:497] + '...'
            
            cleaned_data[field] = value.strip()
    
    return cleaned_data

def calculate_overall_confidence(confidence_scores: Dict[str, float]) -> float:
    """
    Calculate overall confidence score for the entire extraction.
    """
    if not confidence_scores:
        return 0.0
    
    # Weight different sections based on importance for plagiarism detection
    weights = {
        'title': 0.1,
        'authors': 0.1,
        'abstract': 0.2,
        'introduction': 0.15,
        'methodology': 0.15,
        'results': 0.15,
        'conclusion': 0.1,
        'references': 0.05
    }
    
    weighted_sum = 0.0
    total_weight = 0.0
    
    for section, weight in weights.items():
        if section in confidence_scores:
            weighted_sum += confidence_scores[section] * weight
            total_weight += weight
    
    if total_weight > 0:
        return weighted_sum / total_weight
    
    return sum(confidence_scores.values()) / len(confidence_scores) if confidence_scores else 0.0

def analyze_content_with_groq_enterprise(text: str) -> ExtractionResult:
    """
    Enterprise-grade content analysis with comprehensive error handling and confidence scoring.
    """
    start_time = datetime.now()
    errors = []
    
    try:
        logger.info("Starting enterprise-grade content analysis")
        
        # Step 1: Advanced document section extraction
        logger.info("Step 1: Advanced section extraction")
        extracted_sections, confidence_scores, extraction_errors = extract_document_sections_advanced(text)
        errors.extend(extraction_errors)
        
        # Step 2: AI-powered analysis and enhancement
        logger.info("Step 2: AI-powered analysis")
        enhanced_sections, enhanced_confidence, ai_errors = analyze_extracted_content_with_groq_advanced(extracted_sections, confidence_scores)
        errors.extend(ai_errors)
        
        # Calculate overall confidence
        overall_confidence = calculate_overall_confidence(enhanced_confidence)
        
        # Calculate word count
        word_count = len(text.split())
        
        # Calculate processing time
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Determine extraction method
        extraction_method = "AI-Enhanced" if overall_confidence > 0.7 else "Traditional-Fallback"
        
        logger.info(f"Analysis completed in {processing_time:.2f} seconds with {overall_confidence:.2f} confidence")
        
        return ExtractionResult(
            title=enhanced_sections.get('title', 'Untitled'),
            authors=enhanced_sections.get('authors', 'Unknown'),
            abstract=enhanced_sections.get('abstract', 'No abstract found'),
            introduction=enhanced_sections.get('introduction', ''),
            methodology=enhanced_sections.get('methodology', ''),
            results=enhanced_sections.get('results', ''),
            conclusion=enhanced_sections.get('conclusion', ''),
            references=enhanced_sections.get('references', ''),
            content_type=enhanced_sections.get('content_type', 'RESEARCH_PAPER'),
            confidence_score=overall_confidence,
            extraction_method=extraction_method,
            error_log=errors,
            processing_time=processing_time,
            word_count=word_count,
            section_confidence=enhanced_confidence
        )
        
    except Exception as e:
        error_msg = f"Critical enterprise analysis error: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        
        # Return minimal result with error information
        return ExtractionResult(
            title="Error in Analysis",
            authors="Unknown",
            abstract="Analysis failed",
            introduction="",
            methodology="",
            results="",
            conclusion="",
            references="",
            content_type="RESEARCH_PAPER",
            confidence_score=0.0,
            extraction_method="Error-Fallback",
            error_log=errors,
            processing_time=(datetime.now() - start_time).total_seconds(),
            word_count=0,
            section_confidence={}
        )

def upload_academic_content(request):
    """
    Enterprise-grade file upload handler with comprehensive error handling and confidence reporting.
    """
    message = ''
    error = False
    title = ''
    confidence_info = {}
    
    if request.method == 'POST' and request.FILES.get('file'):
        form = ResearchPaperForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Save the file to the database
                paper = form.save()
                file_path = paper.file.path
                
                logger.info(f"Processing file: {file_path}")
                
                # Extract text from the uploaded file
                extracted_text, extraction_errors = extract_text_from_file(file_path)
                
                if not extracted_text:
                    message = "Could not extract text from the uploaded file."
                    error = True
                    return render(request, 'upload.html', {
                        'form': form, 
                        'message': message, 
                        'error': error, 
                        'title': title
                    })
                
                # Perform enterprise-grade content analysis
                extraction_result = analyze_content_with_groq_enterprise(extracted_text)
                
                # Validate and save extracted content into the database
                paper.title = extraction_result.title if extraction_result.title and extraction_result.title.strip() else 'Untitled'
                paper.authors = extraction_result.authors if extraction_result.authors and extraction_result.authors.strip() else 'Unknown'
                paper.abstract = extraction_result.abstract if extraction_result.abstract and extraction_result.abstract.strip() else ''
                paper.introduction = extraction_result.introduction if extraction_result.introduction and extraction_result.introduction.strip() else ''
                paper.methodology = extraction_result.methodology if extraction_result.methodology and extraction_result.methodology.strip() else ''
                paper.results = extraction_result.results if extraction_result.results and extraction_result.results.strip() else ''
                paper.conclusion = extraction_result.conclusion if extraction_result.conclusion and extraction_result.conclusion.strip() else ''
                paper.references = extraction_result.references if extraction_result.references and extraction_result.references.strip() else ''
                paper.content_type = extraction_result.content_type if extraction_result.content_type else 'RESEARCH_PAPER'
                
                # Ensure data integrity before saving
                try:
                    paper.save()
                    logger.info(f"Successfully saved paper to database: {paper.title}")
                except Exception as db_error:
                    logger.error(f"Database save error: {str(db_error)}")
                    # Try to save with minimal data if full save fails
                    paper.title = 'Untitled'
                    paper.authors = 'Unknown'
                    paper.save()
                
                # Prepare confidence information for display
                confidence_info = {
                    'overall_confidence': f"{extraction_result.confidence_score:.1%}",
                    'extraction_method': extraction_result.extraction_method,
                    'processing_time': f"{extraction_result.processing_time:.2f}s",
                    'word_count': extraction_result.word_count,
                    'section_confidence': extraction_result.section_confidence
                }
                
                message = f"File uploaded and processed successfully! Title: {extraction_result.title}"
                title = extraction_result.title
                
                # Log any errors for monitoring
                if extraction_result.error_log:
                    logger.warning(f"Processing completed with {len(extraction_result.error_log)} warnings")
                    for error_msg in extraction_result.error_log:
                        logger.warning(error_msg)
                
                # Pass all extracted data to template
                context = {
                    'form': form, 
                    'message': message, 
                    'error': error, 
                    'title': extraction_result.title,
                    'content_type': extraction_result.content_type,
                    'extracted_date': paper.extracted_date,
                    'abstract': extraction_result.abstract,
                    'introduction': extraction_result.introduction,
                    'methodology': extraction_result.methodology,
                    'results': extraction_result.results,
                    'conclusion': extraction_result.conclusion,
                    'references': extraction_result.references,
                    'confidence_info': confidence_info
                }
                return render(request, 'upload.html', context)
                
            except ProductionError as e:
                message = f"Production error: {str(e)}"
                error = True
                logger.error(f"Production error in upload: {str(e)}")
            except Exception as e:
                message = f"Unexpected error: {str(e)}"
                error = True
                logger.error(f"Unexpected error in upload: {str(e)}")
                logger.error(traceback.format_exc())
        else:
            message = "There was an issue with the file upload."
            error = True
    else:
        form = ResearchPaperForm()
    
    return render(request, 'upload.html', {
        'form': form, 
        'message': message, 
        'error': error, 
        'title': title,
        'confidence_info': confidence_info
    })

class AcademicContentList(ListView):
    model = ResearchPaper
    template_name = 'academic_content_list.html'
    context_object_name = 'academic_contents'
